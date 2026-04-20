# Here is a **line‑by‑line explanation** of the Microsoft Word (`.docx`) extraction code using `python-docx`, presented as bullet points.

# - `from docx import Document`
#   - Imports the `Document` class from the `python-docx` library.
#   - This class is used to open, read, and manipulate Word `.docx` files.

# - `doc = Document('C:\\Personal\\2024\\Learning\\Generative AI\\RAG\\27_Context_Engineering\\2_RAG\\1_Document_Processing\\Data\\employee_report.docx')`
#   - Creates a `Document` object by opening the Word file located at the given path.
#   - The double backslashes (`\\`) are escape sequences for a single backslash in a regular string.

# - `# ── Method 1: Read all paragraphs with style info ─────────────────────────────`
#   - A comment indicating the start of the first method (extracting paragraphs with their styles).

# - `print("=" * 60)`
#   - Prints a line of 60 equals signs as a visual separator.

# - `print("METHOD 1: Paragraphs with Headings & Styles")`
#   - Prints a descriptive header for the first method.

# - `print("=" * 60)`
#   - Prints another separator line.

# - `for para in doc.paragraphs:`
#   - Iterates over all paragraphs in the document. `doc.paragraphs` returns a list of `Paragraph` objects.

# - `if not para.text.strip():`
#   - Checks if the paragraph contains only whitespace or is empty after stripping.
#   - If true, `continue` skips this paragraph (no content to print).

# - `# Fix: safely get style name, default to 'Normal' if None`
#   - A comment explaining that the style attribute might be `None`, so we provide a fallback.

# - `style = para.style.name if para.style is not None else 'Normal'`
#   - Tries to get the style name of the paragraph (e.g., `'Heading 1'`, `'Normal'`, `'List Paragraph'`).
#   - If `para.style` is `None`, defaults to `'Normal'`.

# - `if style == 'Heading 1':`
#   - Checks if the paragraph style is `'Heading 1'`.

# - `print(f"\n{'#' * 60}")`
#   - Prints a newline followed by 60 hash (`#`) characters to visually highlight a top‑level heading.

# - `print(f"  H1: {para.text}")`
#   - Prints the heading level (`H1:`) and the paragraph text.

# - `print(f"{'#' * 60}")`
#   - Prints another line of 60 hashes to close the heading block.

# - `elif style == 'Heading 2':`
#   - Otherwise, if the style is `'Heading 2'`.

# - `print(f"\n  >> H2: {para.text}")`
#   - Prints a newline, then `>> H2:` followed by the heading text.

# - `print(f"  {'-' * 40}")`
#   - Prints 40 hyphens indented by two spaces as a visual separator.

# - `elif style == 'Heading 3':`
#   - Otherwise, if the style is `'Heading 3'`.

# - `print(f"\n    > H3: {para.text}")`
#   - Prints a newline, then indented `> H3:` and the heading text.

# - `elif 'List' in style:`
#   - Otherwise, if the style name contains the word `'List'` (e.g., `'List Paragraph'`).

# - `print(f"    • {para.text}")`
#   - Prints a bullet symbol (`•`) indented, followed by the paragraph text.

# - `else:`
#   - For any other style (normal text, body text, etc.).

# - `formatted = ""`
#   - Initialises an empty string to hold the formatted output for this paragraph.

# - `for run in para.runs:`
#   - Iterates over each `Run` object within the paragraph. A run is a contiguous sequence of characters with the same formatting (bold, italic, etc.).

# - `if run.bold:`
#   - Checks if the run has bold formatting.

# - `formatted += f"[BOLD: {run.text}]"`
#   - Appends the run’s text wrapped in `[BOLD: ...]` to indicate bold text.

# - `elif run.italic:`
#   - Otherwise, checks if the run has italic formatting.

# - `formatted += f"[ITALIC: {run.text}]"`
#   - Appends the run’s text wrapped in `[ITALIC: ...]` to indicate italic text.

# - `else:`
#   - If the run has neither bold nor italic.

# - `formatted += run.text`
#   - Appends the run’s text as is (no special markers).

# - `print(f"    {formatted}")`
#   - Prints the fully formatted paragraph text, indented by four spaces.

# - `# ── Method 2: Read tables preserving structure ────────────────────────────────`
#   - A comment indicating the start of the second method (extracting tables).

# - `print("\n" + "=" * 60)`
#   - Prints a newline followed by a separator line.

# - `print("METHOD 2: Tables with Structure")`
#   - Prints the header for the second method.

# - `print("=" * 60)`
#   - Prints another separator.

# - `for t_idx, table in enumerate(doc.tables, 1):`
#   - Iterates over all tables in the document. `doc.tables` returns a list of `Table` objects.
#   - `enumerate(..., 1)` provides a table index starting from 1.

# - `print(f"\nTable {t_idx}: {len(table.rows)} rows x {len(table.columns)} cols")`
#   - Prints the table number, number of rows, and number of columns.

# - `print("-" * 60)`
#   - Prints a line of 60 hyphens.

# - `for r_idx, row in enumerate(table.rows):`
#   - Iterates over each row in the current table. `table.rows` returns a list of `Row` objects.

# - `cells = [cell.text.strip() for cell in row.cells]`
#   - Creates a list by iterating over each cell in the row, extracting the text, and stripping whitespace.

# - `label = "HEADER" if r_idx == 0 else f"Row {r_idx:>2}"`
#   - If the row index is 0, assigns the label `"HEADER"` (first row is often a header). Otherwise, assigns `"Row X"` where `X` is right‑aligned to 2 characters.

# - `print(f"  {label}: {' | '.join(cells)}")`
#   - Prints the label, a colon, and the cell values joined by `" | "` (space‑pipe‑space), all indented by two spaces.

# - `# ── Method 3: Full document structure as dict ─────────────────────────────────`
#   - A comment indicating the start of the third method (building a structured dictionary).

# - `print("\n" + "=" * 60)`
#   - Prints a newline and a separator.

# - `print("METHOD 3: Full Document Structure as Dictionary")`
#   - Prints the header for the third method.

# - `print("=" * 60)`
#   - Prints another separator.

# - `structure = {"headings": [], "paragraphs": [], "tables": [], "lists": []}`
#   - Initialises a dictionary with four empty lists to store different document elements.

# - `for para in doc.paragraphs:`
#   - Iterates over all paragraphs again (similar to Method 1, but now building the dictionary).

# - `if not para.text.strip():`
#   - Skips empty or whitespace‑only paragraphs.

# - `# Fix: safely get style name here too`
#   - A comment repeating the safety check for style name.

# - `style = para.style.name if para.style is not None else 'Normal'`
#   - Retrieves the style name, defaulting to `'Normal'`.

# - `if 'Heading' in style:`
#   - Checks if the style name contains the word `'Heading'`.

# - `structure["headings"].append({"level": style, "text": para.text})`
#   - Appends a dictionary with the style name and paragraph text to the `"headings"` list.

# - `elif 'List' in style:`
#   - Otherwise, if the style contains `'List'`.

# - `structure["lists"].append(para.text)`
#   - Appends the paragraph text to the `"lists"` list.

# - `else:`
#   - For any other style.

# - `structure["paragraphs"].append(`
#   - Appends a truncated version of the paragraph text to the `"paragraphs"` list.

# - `para.text[:80] + "..." if len(para.text) > 80 else para.text`
#   - If the paragraph text is longer than 80 characters, it takes the first 80 characters and adds `"..."`. Otherwise, keeps the full text.

# - `)`
#   - Closes the `append` call.

# - `for table in doc.tables:`
#   - Iterates over all tables again.

# - `rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]`
#   - Creates a nested list: each inner list contains the stripped text of all cells in a row.

# - `structure["tables"].append({`
#   - Appends a dictionary to the `"tables"` list.

# - `"rows": len(table.rows),`
#   - Stores the number of rows.

# - `"cols": len(table.columns),`
#   - Stores the number of columns.

# - `"data": rows`
#   - Stores the actual row data (the nested list).

# - `})`
#   - Closes the dictionary and the `append` call.

# - `print(f"  Headings  : {len(structure['headings'])}")`
#   - Prints the count of headings found.

# - `print(f"  Paragraphs: {len(structure['paragraphs'])}")`
#   - Prints the count of paragraphs found.

# - `print(f"  Tables    : {len(structure['tables'])}")`
#   - Prints the number of tables.

# - `print(f"  List items: {len(structure['lists'])}")`
#   - Prints the number of list items.

# - `print("\n  Headings found:")`
#   - Prints a newline and a sub‑header.

# - `for h in structure["headings"]:`
#   - Iterates over the headings dictionary list.

# - `print(f"    [{h['level']}] {h['text']}")`
#   - Prints each heading with its style level and text, indented.

# - `print("\n  Table data:")`
#   - Prints a newline and a sub‑header.

# - `for i, tbl in enumerate(structure["tables"], 1):`
#   - Iterates over the tables with a 1‑based index.

# - `print(f"    Table {i} ({tbl['rows']}x{tbl['cols']}):")`
#   - Prints the table number and its dimensions.

# - `for row in tbl["data"]:`
#   - Iterates over the rows of the current table.

# - `print(f"      {row}")`
#   - Prints each row (as a list) indented further.

from docx import Document

doc = Document('D:\GenAI Content\\AI code\\4_RAG_Indexing\\1_Document_Processing\\Data\\employee_report.docx')

# ── Method 1: Read all paragraphs with style info ─────────────────────────────
print("=" * 60)
print("METHOD 1: Paragraphs with Headings & Styles")
print("=" * 60)

for para in doc.paragraphs:
    if not para.text.strip():
        continue

    # Fix: safely get style name, default to 'Normal' if None
    style = para.style.name if para.style is not None else 'Normal'

    if style == 'Heading 1':
        print(f"\n{'#' * 60}")
        print(f"  H1: {para.text}")
        print(f"{'#' * 60}")
    elif style == 'Heading 2':
        print(f"\n  >> H2: {para.text}")
        print(f"  {'-' * 40}")
    elif style == 'Heading 3':
        print(f"\n    > H3: {para.text}")
    elif 'List' in style:
        print(f"    • {para.text}")
    else:
        formatted = ""
        for run in para.runs:
            if run.bold:
                formatted += f"[BOLD: {run.text}]"
            elif run.italic:
                formatted += f"[ITALIC: {run.text}]"
            else:
                formatted += run.text
        print(f"    {formatted}")

# ── Method 2: Read tables preserving structure ────────────────────────────────
print("\n" + "=" * 60)
print("METHOD 2: Tables with Structure")
print("=" * 60)

for t_idx, table in enumerate(doc.tables, 1):
    print(f"\nTable {t_idx}: {len(table.rows)} rows x {len(table.columns)} cols")
    print("-" * 60)

    for r_idx, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        label = "HEADER" if r_idx == 0 else f"Row {r_idx:>2}"
        print(f"  {label}: {' | '.join(cells)}")

# ── Method 3: Full document structure as dict ─────────────────────────────────
print("\n" + "=" * 60)
print("METHOD 3: Full Document Structure as Dictionary")
print("=" * 60)

structure = {"headings": [], "paragraphs": [], "tables": [], "lists": []}

for para in doc.paragraphs:
    if not para.text.strip():
        continue

    # Fix: safely get style name here too
    style = para.style.name if para.style is not None else 'Normal'

    if 'Heading' in style:
        structure["headings"].append({"level": style, "text": para.text})
    elif 'List' in style:
        structure["lists"].append(para.text)
    else:
        structure["paragraphs"].append(
            para.text[:80] + "..." if len(para.text) > 80 else para.text
        )

for table in doc.tables:
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    structure["tables"].append({
        "rows": len(table.rows),
        "cols": len(table.columns),
        "data": rows
    })

print(f"  Headings  : {len(structure['headings'])}")
print(f"  Paragraphs: {len(structure['paragraphs'])}")
print(f"  Tables    : {len(structure['tables'])}")
print(f"  List items: {len(structure['lists'])}")

print("\n  Headings found:")
for h in structure["headings"]:
    print(f"    [{h['level']}] {h['text']}")

print("\n  Table data:")
for i, tbl in enumerate(structure["tables"], 1):
    print(f"    Table {i} ({tbl['rows']}x{tbl['cols']}):")
    for row in tbl["data"]:
        print(f"      {row}")